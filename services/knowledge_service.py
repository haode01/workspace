"""知识库业务逻辑层 —— 文档解析、切分、索引、搜索"""

import os
import re
import hashlib
import logging
from typing import List, Tuple

from data.database import Database
from ai.rag_engine import RAGEngine

logger = logging.getLogger(__name__)


class KnowledgeService:
    def __init__(self, db: Database, rag_engine: RAGEngine):
        self.db = db
        self.rag = rag_engine

    # ────────────────── 编码检测读取 ──────────────────
    @staticmethod
    def _read_text_file(filepath: str) -> str:
        """智能编码读取: UTF-8 → chardet 检测 → GBK 兜底"""
        raw = b""
        try:
            with open(filepath, "rb") as f:
                raw = f.read()
        except Exception as e:
            logger.warning("[文本转换] 无法读取文件: %s", e)
            return ""
        if not raw:
            return ""
        # 1) 尝试 UTF-8 (含 BOM)
        for enc in ("utf-8-sig", "utf-8"):
            try:
                return raw.decode(enc)
            except (UnicodeDecodeError, ValueError):
                pass
        # 2) chardet 自动检测
        try:
            import chardet
            det = chardet.detect(raw)
            if det and det.get("encoding"):
                try:
                    text = raw.decode(det["encoding"])
                    logger.info("[文本转换] chardet检测编码: %s (confidence=%.2f)",
                                det["encoding"], det.get("confidence", 0))
                    return text
                except (UnicodeDecodeError, LookupError):
                    pass
        except ImportError:
            pass
        # 3) 常见中文编码兜底
        for enc in ("gb18030", "gbk", "gb2312", "big5", "latin-1"):
            try:
                return raw.decode(enc)
            except (UnicodeDecodeError, ValueError):
                pass
        # 最终 fallback
        return raw.decode("utf-8", errors="replace")

    # ────────────────── 文件解析 (统一转 Markdown) ──────────────────
    def parse_file(self, filepath: str) -> str:
        """解析各种格式文件, 统一输出 Markdown 文本"""
        ext = os.path.splitext(filepath)[1].lower()
        logger.info("[文本转换] 开始解析文件: %s (类型: %s)", filepath, ext)

        raw = ""
        if ext == ".pdf":
            raw = self._parse_pdf(filepath)
        elif ext == ".docx":
            raw = self._parse_docx(filepath)
        elif ext in (".html", ".htm"):
            raw = self._parse_html(filepath)
        elif ext in (".txt", ".md", ".csv", ".json", ".py", ".js", ".java", ".c", ".cpp"):
            raw = self._read_text_file(filepath)
        else:
            # 尝试当文本读取
            raw = self._read_text_file(filepath)
            if not raw:
                logger.warning("[文本转换] 无法读取文件: %s", ext)
                return ""

        if self._has_garbled_text(raw):
            logger.warning("[文本转换] 检测到可能的乱码内容! 文件: %s (前100字: %s)",
                           filepath, raw[:100].replace('\n', ' '))

        md = self._to_markdown(raw, ext)
        logger.info("[文本转换] 转换完成, 原始长度=%d, MD长度=%d", len(raw), len(md))
        return md

    @staticmethod
    def _has_garbled_text(text: str) -> bool:
        """检测文本是否可能是乱码 (大量非常见 Unicode 字符)"""
        if not text:
            return False
        suspicious = 0
        total = 0
        for ch in text:
            if ch.isspace() or ch in ('\n', '\r', '\t'):
                continue
            total += 1
            cp = ord(ch)
            # 常见范围: ASCII(20-7E), 中文(4E00-9FFF), 中文标点(3000-303F, FF00-FFEF)
            if not ((0x20 <= cp <= 0x7E) or (0x4E00 <= cp <= 0x9FFF)
                    or (0x3000 <= cp <= 0x303F) or (0xFF00 <= cp <= 0xFFEF)
                    or (0x2000 <= cp <= 0x206F)):
                suspicious += 1
        if total == 0:
            return False
        ratio = suspicious / total
        return ratio > 0.15  # 超过15%的字符不在常见范围内

    @staticmethod
    def _parse_pdf(filepath: str) -> str:
        text = ""
        # 优先 pdfplumber (中文支持更好)
        try:
            import pdfplumber
            with pdfplumber.open(filepath) as pdf:
                pages = [p.extract_text() or "" for p in pdf.pages]
                text = "\n\n".join(pages)
                logger.info("[文本转换] pdfplumber解析PDF: %d页", len(pdf.pages))
                if text.strip():
                    return text
        except ImportError:
            pass
        except Exception as e:
            logger.warning("[文本转换] pdfplumber解析失败, 尝试PyPDF2: %s", e)
        # 回退 PyPDF2
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(filepath)
            logger.info("[文本转换] PyPDF2解析PDF: %d页", len(reader.pages))
            text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
            return text
        except ImportError:
            logger.error("[文本转换] PyPDF2 未安装, 请执行: pip install PyPDF2")
            return ""
        except Exception as e:
            logger.error("[文本转换] PDF解析失败: %s", e)
            return ""

    @staticmethod
    def _parse_docx(filepath: str) -> str:
        try:
            from docx import Document as DocxDocument
            from docx.oxml.ns import qn
        except ImportError:
            logger.error("[文本转换] python-docx 未安装, 请执行: pip install python-docx")
            return ""
        try:
            import base64
            doc = DocxDocument(filepath)
            parts: list[str] = []
            img_count = 0

            # 构建 rId → 图片数据 映射
            rels = {}
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    try:
                        blob = rel.target_part.blob
                        ct = rel.target_part.content_type or "image/png"
                        b64 = base64.b64encode(blob).decode("ascii")
                        rels[rel.rId] = f"data:{ct};base64,{b64}"
                    except Exception:
                        pass

            def _extract_images_from_paragraph(para) -> list[str]:
                """从段落 XML 中提取内联图片"""
                imgs = []
                for drawing in para._element.findall(f".//{qn('w:drawing')}"):
                    for blip in drawing.findall(f".//{qn('a:blip')}"):
                        rId = blip.get(qn("r:embed"))
                        if rId and rId in rels:
                            imgs.append(rels[rId])
                return imgs

            # 预构建 element → 对象映射, 避免 O(n²) 查找
            para_map = {p._element: p for p in doc.paragraphs}
            table_map = {t._element: t for t in doc.tables}

            for element in doc.element.body:
                tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

                if tag == "p":
                    p = para_map.get(element)
                    if not p:
                        continue
                    text = p.text.strip()
                    imgs = _extract_images_from_paragraph(p)
                    style_name = (p.style.name or "").lower() if p.style else ""
                    if text:
                        if "heading 1" in style_name:
                            parts.append(f"# {text}")
                        elif "heading 2" in style_name:
                            parts.append(f"## {text}")
                        elif "heading 3" in style_name:
                            parts.append(f"### {text}")
                        else:
                            parts.append(text)
                    for data_uri in imgs:
                        img_count += 1
                        parts.append(f"![图片{img_count}]({data_uri})")

                elif tag == "tbl":
                    table = table_map.get(element)
                    if not table:
                        continue
                    rows_md = []
                    for ri, row in enumerate(table.rows):
                        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                        rows_md.append("| " + " | ".join(cells) + " |")
                        if ri == 0:
                            rows_md.append("| " + " | ".join("---" for _ in cells) + " |")
                    if rows_md:
                        parts.append("\n".join(rows_md))

            logger.info("[文本转换] DOCX段落数: %d, 表格数: %d, 图片数: %d",
                        len(doc.paragraphs), len(doc.tables), img_count)
            return "\n\n".join(parts)
        except Exception as e:
            logger.error("[文本转换] DOCX解析失败: %s", e)
            return ""

    @classmethod
    def _parse_html(cls, filepath: str) -> str:
        html = cls._read_text_file(filepath)
        # 简单去标签
        text = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL)
        text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL)
        text = re.sub(r"<[^>]+>", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    @staticmethod
    def _to_markdown(text: str, ext: str) -> str:
        """将原始文本转换为 Markdown 格式存储"""
        if not text.strip():
            return ""
        if ext == ".md":
            return text
        if ext in (".py", ".js", ".java", ".c", ".cpp", ".go", ".rs", ".ts"):
            lang = ext.lstrip(".")
            return f"```{lang}\n{text}\n```"
        if ext == ".csv":
            lines = text.strip().split("\n")
            if len(lines) >= 2:
                header = lines[0].split(",")
                md = "| " + " | ".join(h.strip() for h in header) + " |\n"
                md += "| " + " | ".join("---" for _ in header) + " |\n"
                for line in lines[1:]:
                    cols = line.split(",")
                    md += "| " + " | ".join(c.strip() for c in cols) + " |\n"
                return md
            return text
        if ext == ".json":
            return f"```json\n{text}\n```"
        # PDF / DOCX / HTML / TXT → 段落化
        paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
        return "\n\n".join(paragraphs)

    # ────────────────── 文本切分 ──────────────────
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 100) -> List[str]:
        logger.info("[文本转换] 开始切分文本, 总长度: %d, chunk_size: %d, overlap: %d", len(text), chunk_size, overlap)
        chunks = []
        start = 0
        while start < len(text):
            chunk = text[start: start + chunk_size].strip()
            if chunk:
                chunks.append(chunk)
            start += chunk_size - overlap
        logger.info("[文本转换] 切分完成, 生成 %d 个文本块", len(chunks))
        return chunks

    # ────────────────── 图片剥离 (用于切分/embedding) ──────────────────
    _IMG_RE = re.compile(r'!\[([^\]]*)\]\(data:[^)]+\)')

    @classmethod
    def _strip_images(cls, text: str) -> str:
        """将 base64 图片替换为轻量占位符, 避免切块/embedding 爆炸"""
        return cls._IMG_RE.sub(r'[图片: \1]', text)

    # ────────────────── 文档导入 ──────────────────
    def add_document(self, filepath: str, ai_client=None, category: str = "", original_filename: str = "") -> int:
        filename = original_filename or os.path.basename(filepath)
        ext = os.path.splitext(filepath)[1].lower().lstrip(".")
        logger.info("[存储] 开始导入文档: %s (类型: %s, 分类: %s)", filename, ext, category)
        content = self.parse_file(filepath)
        if not content.strip():
            logger.warning("[存储] 文档内容为空, 跳过导入: %s", filename)
            return -1

        # 先用原文存储, 立即返回 doc_id
        doc_id = self.db.add_document(filename, content, ext, category)
        logger.info("[存储] 文档已存入DB(原文), doc_id=%d, 长度=%d", doc_id, len(content))

        # 先用原文建立索引 (快速可搜索)
        self._index_document(doc_id, content, filename, category, ai_client)

        # 后台异步: AI 排版 → 更新内容 → 重建索引
        if ai_client and ai_client.is_configured():
            self._async_format(doc_id, content, filename, category, ai_client)

        return doc_id

    def update_text_document(self, doc_id: int, title: str, content: str, category: str = "", ai_client=None) -> bool:
        """更新文本文档并重建索引"""
        if not content.strip():
            logger.warning("[存储] 更新文本失败: 内容为空, doc_id=%d", doc_id)
            return False

        ext = "md"
        filename = title.strip() or "未命名文档"
        if not filename.endswith(f".{ext}"):
            filename = f"{filename}.{ext}"

        doc = self.db.get_document_by_id(doc_id)
        if not doc:
            logger.warning("[存储] 更新文本失败: 文档不存在, doc_id=%d", doc_id)
            return False

        self.db.update_document(doc_id, filename, content, category)
        self.db.delete_chunks_by_doc(doc_id)
        self.rag.remove_doc_chunks(doc_id)
        self._index_document(doc_id, content, filename, category, ai_client)
        logger.info("[存储] 文本文档更新完成: doc_id=%d, filename='%s', category='%s'", doc_id, filename, category)
        return True

    def _index_document(self, doc_id: int, content: str, filename: str, category: str, ai_client=None):
        """为文档内容建立 chunk + embedding 索引"""
        search_text = self._strip_images(content)
        chunks = self.chunk_text(search_text)

        if ai_client and ai_client.is_configured() and ai_client.embedding_model:
            try:
                logger.info("[存储] 使用AI生成embedding, chunks数=%d", len(chunks))
                embeddings = ai_client.get_embeddings(chunks)
                logger.info("[存储] AI embedding生成成功")
            except Exception as e:
                logger.warning("[存储] AI embedding失败, 回退到哈希方案: %s", e)
                embeddings = self._fallback_embeddings(chunks)
        else:
            embeddings = self._fallback_embeddings(chunks)

        chunk_dicts = []
        for i, (c, _e) in enumerate(zip(chunks, embeddings)):
            self.db.add_chunk(doc_id, c, i)
            chunk_dicts.append({
                "id": i, "doc_id": doc_id, "content": c,
                "filename": filename, "category": category,
            })
        self.rag.add_chunks(chunk_dicts, embeddings)
        logger.info("[存储] 索引建立完成: doc_id=%d, chunks=%d", doc_id, len(chunk_dicts))

    def _async_format(self, doc_id: int, raw_content: str, filename: str, category: str, ai_client):
        """后台线程: AI 排版 → 更新 DB → 重建索引"""
        import threading

        def _task():
            try:
                logger.info("[异步排版] 开始: doc_id=%d, 长度=%d", doc_id, len(raw_content))
                formatted = self._format_to_markdown(raw_content, ai_client)
                if formatted == raw_content:
                    logger.info("[异步排版] 内容未变化, 跳过更新: doc_id=%d", doc_id)
                    return
                # 更新 DB 内容
                self.db.update_document_content(doc_id, formatted)
                # 删除旧 chunks, 重建索引
                self.db.delete_chunks_by_doc(doc_id)
                self.rag.remove_doc_chunks(doc_id)
                self._index_document(doc_id, formatted, filename, category, ai_client)
                logger.info("[异步排版] 完成: doc_id=%d", doc_id)
            except Exception as e:
                logger.error("[异步排版] 失败: doc_id=%d, %s", doc_id, e)

        t = threading.Thread(target=_task, daemon=True)
        t.start()
        logger.info("[异步排版] 后台任务已启动: doc_id=%d", doc_id)

    _MD_FORMAT_SYSTEM = (
        "你是一个纯排版工具。你的唯一任务是将用户的原始文本转换为 Markdown 格式。\n\n"
        "【铁律 — 必须严格遵守】\n"
        "- 禁止添加任何原文中不存在的内容、总结、解释、评论、前言、结语\n"
        "- 禁止删除或省略原文的任何文字，每一句话都必须保留\n"
        "- 禁止改写、润色、同义替换原文措辞\n"
        "- 你只能做：加标题标记(#)、加列表标记(-)、加粗(**)、加代码块(```)、加表格、分段换行\n\n"
        "【排版规则】\n"
        "1. 识别标题层级，添加 #、##、### 标记\n"
        "2. 合理分段，避免超长段落\n"
        "3. 枚举内容用 - 或 1. 列表\n"
        "4. 对比内容用表格\n"
        "5. 代码/命令/路径用代码块包裹\n"
        "6. 重要概念加粗\n\n"
        "【特殊标记】\n"
        "- 文本中的 __IMG_PLACEHOLDER_数字__ 是图片占位符，必须原样保留在原位置，不要删除或修改\n\n"
        "【输出要求】\n"
        "- 直接输出 Markdown，不要任何包裹或解释\n"
        "- 不要用 ```markdown 包裹输出"
    )

    # 中文 1 token ≈ 1.5 字符, 英文 1 token ≈ 4 字符; 取保守值 2 字符/token
    # 50k 字符 ≈ 25k tokens 输入, 预留 system(~1k) + 输出(最多与输入等长)
    # 总计 ≈ 25k + 1k + 25k = 51k tokens, 在 131k 限制内安全
    _MD_CHUNK_CHARS = 50000
    _MD_MAX_WORKERS = 4       # 并行调用 AI 的线程数

    @staticmethod
    def _already_markdown(text: str) -> bool:
        """检测文本是否已经是结构化 Markdown, 无需再格式化
        必须同时有标题(#)和其他标记才算已格式化, 仅有表格/列表不算"""
        lines = text.strip().split('\n')
        if len(lines) < 3:
            return False
        headings = 0
        other_markers = 0
        for line in lines[:100]:
            s = line.strip()
            if s.startswith('#'):
                headings += 1
            elif s.startswith(('- ', '* ', '1. ', '2. ', '3. ')):
                other_markers += 1
            elif s.startswith('```'):
                other_markers += 1
            elif '**' in s:
                other_markers += 0.5
            # 表格行不计入 (原始粘贴的表格也含 |)
        # 至少要有 2 个标题 + 一定比例的其他标记
        if headings < 2:
            return False
        total = headings * 2 + other_markers
        ratio = total / min(len(lines), 100)
        return ratio > 0.15

    # 匹配 Markdown 图片语法 (含 base64 和普通 URL)
    _IMG_PLACEHOLDER_RE = re.compile(r'(!\[[^\]]*\]\([^)]+\))')

    def _format_to_markdown(self, raw_text: str, ai_client) -> str:
        """使用 AI 将原始文本结构化为高质量 Markdown (自动分段 + 并行处理)"""
        text = raw_text.strip()
        if not text:
            return raw_text

        # 已经是 Markdown → 跳过
        if self._already_markdown(text):
            logger.info("[存储] 文本已是 Markdown 格式, 跳过 AI 排版 (%d 字符)", len(text))
            return raw_text

        # 剥离图片: 用占位符替换, 排版后还原 (避免图片数据浪费 tokens)
        img_store = []
        def _replace_img(m):
            idx = len(img_store)
            img_store.append(m.group(0))
            return f"__IMG_PLACEHOLDER_{idx}__"
        text = self._IMG_PLACEHOLDER_RE.sub(_replace_img, text)
        if img_store:
            logger.info("[存储] 剥离 %d 张图片后送 AI 排版 (%d → %d 字符)",
                        len(img_store), len(raw_text), len(text))

        # 短文本直接处理
        if len(text) <= self._MD_CHUNK_CHARS:
            result = self._md_format_one(text, ai_client)
            return self._restore_images(result, img_store)

        # 长文本: 按段落边界分段, 并行格式化
        logger.info("[存储] 文本过长 (%d 字符), 分段并行 Markdown 格式化", len(text))
        segments = self._split_for_format(text, self._MD_CHUNK_CHARS)
        total = len(segments)
        logger.info("[存储] 分为 %d 段, 并行度=%d", total, self._MD_MAX_WORKERS)

        from concurrent.futures import ThreadPoolExecutor, as_completed

        parts = [None] * total
        with ThreadPoolExecutor(max_workers=self._MD_MAX_WORKERS) as pool:
            future_map = {
                pool.submit(self._md_format_one, seg, ai_client): idx
                for idx, seg in enumerate(segments)
            }
            for future in as_completed(future_map):
                idx = future_map[future]
                parts[idx] = future.result()
                logger.info("[存储] 格式化完成 %d/%d 段", sum(1 for p in parts if p is not None), total)

        result = "\n\n---\n\n".join(parts)
        logger.info("[存储] 分段格式化完成: %d → %d 字符", len(raw_text), len(result))
        return self._restore_images(result, img_store)

    @staticmethod
    def _restore_images(text: str, img_store: list) -> str:
        """将占位符还原为原始图片"""
        if not img_store:
            return text
        for i, img in enumerate(img_store):
            text = text.replace(f"__IMG_PLACEHOLDER_{i}__", img)
        return text

    def _md_format_one(self, text: str, ai_client) -> str:
        """格式化单段文本为 Markdown"""
        try:
            md = ai_client.generate_text(text, system=self._MD_FORMAT_SYSTEM)
            md = md.strip()
            for prefix in ("```markdown", "```md", "```"):
                if md.startswith(prefix):
                    md = md[len(prefix):].strip()
                    break
            if md.endswith("```"):
                md = md[:-3].strip()
            return md
        except Exception as e:
            logger.warning("[存储] AI Markdown 格式化失败, 使用原文: %s", e)
            return text

    @staticmethod
    def _split_for_format(text: str, max_chars: int) -> list:
        """按段落边界将长文本拆分为不超过 max_chars 的段"""
        paragraphs = re.split(r'\n{2,}', text)
        segments, current = [], ""
        for p in paragraphs:
            if current and len(current) + len(p) + 2 > max_chars:
                segments.append(current.strip())
                current = ""
            current += "\n\n" + p if current else p
            # 单段落超限 → 按行硬切
            while len(current) > max_chars:
                cut = current[:max_chars].rfind('\n')
                if cut < max_chars // 2:
                    cut = max_chars
                segments.append(current[:cut].strip())
                current = current[cut:].strip()
        if current.strip():
            segments.append(current.strip())
        return segments

    def add_text_document(self, title: str, content: str, ai_client=None, category: str = "") -> int:
        """从文本输入创建文档 (非文件上传, AI 自动排版为 Markdown)"""
        if not content.strip():
            logger.warning("[存储] 文本内容为空, 跳过导入")
            return -1
        ext = "md"
        filename = title.strip() or "手动输入"
        if not filename.endswith(f".{ext}"):
            filename = f"{filename}.{ext}"
        logger.info("[存储] 文本输入导入: '%s', 分类=%s, 长度=%d", filename, category, len(content))

        # 先用原文存储, 立即返回
        doc_id = self.db.add_document(filename, content, ext, category)
        logger.info("[存储] 文档已存入DB(原文), doc_id=%d", doc_id)

        self._index_document(doc_id, content, filename, category, ai_client)

        # 后台异步 AI 排版
        if ai_client and ai_client.is_configured():
            self._async_format(doc_id, content, filename, category, ai_client)

        return doc_id

    def get_categories(self) -> list:
        return self.db.get_categories()

    # ────────────────── 搜索 (向量检索 + AI 重排序) ──────────────────
    def search(self, query: str, ai_client=None, top_k: int = 5, category: str = "") -> List[Tuple[dict, float]]:
        logger.info("[查询] 知识库搜索: query='%s', top_k=%d, category='%s'", query[:100], top_k, category)

        use_real_embedding = False
        if ai_client and ai_client.is_configured() and ai_client.embedding_model:
            try:
                query_emb = ai_client.get_embedding(query)
                use_real_embedding = True
                logger.info("[查询] AI embedding生成成功, 维度=%d", len(query_emb))
            except Exception as e:
                logger.warning("[查询] AI embedding失败: %s", e)

        # 分类过滤: 增大召回量, 过滤后再截断
        effective_k = top_k * 3 if category else top_k

        # 无真实embedding → 关键词匹配 (比哈希向量可靠得多)
        if not use_real_embedding:
            logger.info("[查询] 无可用embedding模型, 使用关键词匹配")
            raw = self._enrich_results(self._keyword_search(query, effective_k))
            return self._filter_by_category(raw, category, top_k)

        # 维度检查: 旧索引维度与新模型不匹配时降级到关键词
        if self.rag.dimension and len(query_emb) != self.rag.dimension:
            logger.warning("[查询] 维度不匹配 (query=%d, index=%d), 请重建索引, 降级关键词搜索",
                           len(query_emb), self.rag.dimension)
            raw = self._enrich_results(self._keyword_search(query, effective_k))
            return self._filter_by_category(raw, category, top_k)

        # 哈希向量检测: dim=64 通常是哈希兆底的假embedding, 直接用关键词搜索更可靠
        if self.rag.dimension and self.rag.dimension <= 128 and len(query_emb) > 128:
            logger.warning("[查询] 检测到旧的哈希向量 (dim=%d), 跳过向量搜索, 使用关键词+向量混合",
                           self.rag.dimension)
            raw = self._enrich_results(self._keyword_search(query, effective_k))
            return self._filter_by_category(raw, category, top_k)

        # 第一阶段: 向量粗召回 (取 3x 候选, 阈值 0.5 过滤)
        recall_k = min(effective_k * 3, 30)
        candidates = self.rag.search(query_emb, recall_k, min_score=0.5)
        logger.info("[查询] 粗召回 %d 条候选", len(candidates))

        if not candidates:
            # 向量搜索无结果时降级到关键词搜索
            logger.info("[查询] 向量召回无结果, 降级关键词搜索")
            raw = self._enrich_results(self._keyword_search(query, effective_k))
            return self._filter_by_category(raw, category, top_k)

        # 混合验证: 对向量召回结果进行关键词校验
        query_lower = query.strip().lower()
        tokens = self._tokenize(query_lower)
        reranked = []
        for chunk, vec_score in candidates:
            content_lower = chunk.get("content", "").lower()
            if query_lower in content_lower:
                reranked.append((chunk, vec_score * 1.3))
            elif tokens and sum(1 for t in tokens if t in content_lower) >= max(1, len(tokens) * 0.5):
                # 部分token命中, 保留但不加分
                reranked.append((chunk, vec_score * 0.8))
            else:
                # 关键词完全未命中, 丢弃
                logger.debug("[查询] 向量召回但关键词未命中, 丢弃: score=%.3f, content=%s",
                             vec_score, content_lower[:60])
        reranked.sort(key=lambda x: x[1], reverse=True)
        candidates = reranked

        if not candidates:
            logger.info("[查询] 向量召回经关键词校验后无结果, 降级关键词搜索")
            raw = self._enrich_results(self._keyword_search(query, effective_k))
            return self._filter_by_category(raw, category, top_k)

        # 第二阶段: AI 重排序
        if ai_client and ai_client.is_configured() and len(candidates) > effective_k:
            try:
                results = self._ai_rerank(query, candidates, ai_client, effective_k)
                logger.info("[查询] AI重排序完成, 返回 %d 条", len(results))
                raw = self._enrich_results(results)
                return self._filter_by_category(raw, category, top_k)
            except Exception as e:
                logger.warning("[查询] AI重排序失败, 使用原始排序: %s", e)

        results = candidates[:effective_k]
        logger.info("[查询] 搜索完成, 返回 %d 条结果", len(results))
        raw = self._enrich_results(results)
        return self._filter_by_category(raw, category, top_k)

    def _filter_by_category(self, results: List[Tuple[dict, float]], category: str, top_k: int) -> List[Tuple[dict, float]]:
        """按分类过滤, 支持层级前缀匹配: '技术' 匹配 '技术/前端/React'"""
        if not category:
            return results[:top_k]
        cat = category.strip()
        filtered = []
        for chunk, score in results:
            c = chunk.get("category", "")
            if c == cat or c.startswith(cat + "/"):
                filtered.append((chunk, score))
        logger.info("[查询] 分类过滤 '%s': %d → %d", cat, len(results), len(filtered))
        return filtered[:top_k]

    def _enrich_results(self, results: List[Tuple[dict, float]]) -> List[Tuple[dict, float]]:
        """为搜索结果补充 filename / category (旧数据可能缺失)"""
        doc_cache: dict = {}
        for chunk, score in results:
            if chunk.get("filename"):
                continue
            doc_id = chunk.get("doc_id")
            if doc_id is None:
                continue
            if doc_id not in doc_cache:
                doc_cache[doc_id] = self.db.get_document_by_id(doc_id)
            doc = doc_cache[doc_id]
            if doc:
                chunk["filename"] = doc.get("filename", "")
                chunk["category"] = doc.get("category", "")
        return results

    def _keyword_search(self, query: str, top_k: int) -> List[Tuple[dict, float]]:
        """基于关键词的搜索 — 当无embedding模型时使用
        采用 TF 密度 + 位置加权, 过滤纯路径/偶然提及"""
        if not self.rag.chunks:
            return []
        query_lower = query.strip().lower()
        if not query_lower:
            return []

        # 将查询拆分为有意义的词条 (中文按字/词, 英文按空格+整词)
        tokens = self._tokenize(query_lower)
        if not tokens:
            return []

        scored = []
        for chunk in self.rag.chunks:
            content = chunk.get("content", "")
            content_lower = content.lower()
            if not content_lower:
                continue

            # ── 计算完整查询串出现次数 ──
            full_count = content_lower.count(query_lower)

            if full_count == 0:
                # 拆分token匹配
                hit = sum(1 for t in tokens if t in content_lower)
                if hit == 0:
                    continue
                token_ratio = hit / len(tokens)
                if token_ratio < 0.5:
                    continue
                # 基础分 = token命中率 * 0.6
                score = token_ratio * 0.6
            else:
                # ── TF 密度: 出现次数 / 内容长度 (归一化) ──
                tf_density = full_count / (len(content_lower) / 100.0)  # 每100字符出现次数
                tf_density = min(tf_density, 1.0)  # 上限1.0

                # ── 位置权重: 越早出现越重要 ──
                first_pos = content_lower.index(query_lower)
                position_score = 1.0 - (first_pos / len(content_lower)) * 0.3

                # ── 路径惩罚: 如果关键词只出现在路径中而非正文, 大幅降权 ──
                path_penalty = 1.0
                non_path_hits = 0
                m_start = 0
                while m_start < len(content_lower):
                    idx = content_lower.find(query_lower, m_start)
                    if idx == -1:
                        break
                    # 检查命中位置前后是否是路径字符 (/, \, -, _)
                    before = content_lower[idx - 1] if idx > 0 else ' '
                    after_pos = idx + len(query_lower)
                    after = content_lower[after_pos] if after_pos < len(content_lower) else ' '
                    if before not in ('/', '\\', '-', '_') and after not in ('/', '\\', '-', '_'):
                        non_path_hits += 1
                    m_start = idx + len(query_lower)
                if non_path_hits == 0 and full_count > 0:
                    path_penalty = 0.4  # 纯路径提及, 大幅降权

                # 综合评分
                score = (0.5 + 0.2 * tf_density + 0.15 * position_score + 0.15) * path_penalty

            # 最低分过滤
            if score < 0.3:
                continue
            scored.append((chunk, round(score, 4)))

        scored.sort(key=lambda x: x[1], reverse=True)
        logger.info("[查询] 关键词匹配完成, 返回 %d 条 (总候选 %d)",
                    min(top_k, len(scored)), len(scored))
        return scored[:top_k]

    @staticmethod
    def _tokenize(text: str) -> List[str]:
        """简单分词: 英文按空格, 中文按2-gram, 过滤单字符"""
        tokens = []
        buf = ""
        for ch in text:
            if '\u4e00' <= ch <= '\u9fff':  # 中文字符
                if buf.strip():
                    tokens.append(buf.strip())
                    buf = ""
                tokens.append(ch)
            elif ch in (' ', '\t', '\n', ',', '.', ';', ':', '!', '?', '/', '\\', '_', '-'):
                if buf.strip():
                    tokens.append(buf.strip())
                    buf = ""
            else:
                buf += ch
        if buf.strip():
            tokens.append(buf.strip())
        # 过滤掉长度<=1的纯英文token (单字母太容易误匹配)
        return [t for t in tokens if len(t) > 1 or '\u4e00' <= t <= '\u9fff']

    @staticmethod
    def _ai_rerank(query: str, candidates: List[Tuple[dict, float]],
                   ai_client, top_k: int) -> List[Tuple[dict, float]]:
        """用 AI 对候选结果按相关性重排序"""
        numbered = "\n".join(
            f"[{i}] {c['content'][:200]}" for i, (c, _) in enumerate(candidates)
        )
        prompt = (
            f"查询: {query}\n\n以下是候选文档片段:\n{numbered}\n\n"
            f"请按与查询的相关性从高到低,返回最相关的{top_k}个编号,"
            f"仅返回逗号分隔的数字编号,如: 2,0,5,1,3"
        )
        resp = ai_client.generate_text(prompt, system="你是搜索排序助手,只输出编号,不要解释。")
        # 解析编号
        import re as _re
        ids = [int(x) for x in _re.findall(r"\d+", resp) if int(x) < len(candidates)]
        seen = set()
        unique_ids = []
        for idx in ids:
            if idx not in seen:
                seen.add(idx)
                unique_ids.append(idx)
        # 补充未出现的
        for i in range(len(candidates)):
            if i not in seen and len(unique_ids) < top_k:
                unique_ids.append(i)
        return [(candidates[i][0], candidates[i][1]) for i in unique_ids[:top_k]]

    def get_documents(self) -> list[dict]:
        docs = self.db.get_documents()
        logger.info("[查询] 获取文档列表, 共 %d 篇", len(docs))
        return docs

    def delete_document(self, doc_id: int):
        logger.info("[存储] 删除文档: doc_id=%d", doc_id)
        self.rag.remove_doc_chunks(doc_id)
        self.db.delete_document(doc_id)
        logger.info("[存储] 文档删除完成: doc_id=%d", doc_id)

    # ────────────────── 哈希兜底 embedding ──────────────────
    @staticmethod
    def _fallback_embeddings(texts: List[str], dim: int = 128) -> List[List[float]]:
        """无 AI 配置时使用 SHA-512 哈希生成伪 embedding（仅做功能演示）"""
        result = []
        for text in texts:
            h = hashlib.sha512(text.encode()).digest()
            vec = [float(b) / 255.0 for b in h[:dim]]
            result.append(vec)
        return result
